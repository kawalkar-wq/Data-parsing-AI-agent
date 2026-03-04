import streamlit as st
import pdfplumber
from docx import Document
from groq import Groq
import io
import os

# --- Page Config ---
st.set_page_config(page_title="Oracle Agent v2.0", page_icon="📘", layout="centered")

st.title("📘 Oracle Agent v2.0")
st.subheader("AI-Powered Course Design Preparer")
st.markdown("Upload a PDF and AI will generate a full Course Design Document instantly.")

# --- Hardcoded Groq API Key (free, no credit card) ---
api_key = "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"

# --- File Uploader ---
uploaded_file = st.file_uploader("📂 Upload your PDF", type=["pdf"])

if uploaded_file:
    if st.button("✨ Generate Course Design Document", use_container_width=True):

        with st.spinner("Step 1/3 — Cleaning PDF..."):
            try:
                clean_text = ""
                with pdfplumber.open(uploaded_file) as pdf:
                    for page in pdf.pages:
                        h, w = page.height, page.width
                        safe_zone = (0, h * 0.10, w, h * 0.90)
                        cropped = page.within_bbox(safe_zone)
                        text = cropped.extract_text()
                        if text:
                            clean_text += text + "\n\n"

                if not clean_text.strip():
                    st.error("❌ No text found. The PDF may be scanned or image-based.")
                    st.stop()

            except Exception as e:
                st.error(f"❌ PDF reading error: {e}")
                st.stop()

        with st.spinner("Step 2/3 — AI is generating your Course Design..."):
            try:
                client = Groq(api_key=api_key)
                chat_completion = client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    max_tokens=4096,
                    messages=[
                        {
                            "role": "user",
                            "content": f"""You are an expert Instructional Designer specializing in technical product training.
Your job is to analyze product documentation and design a comprehensive, role-specific
training course for Technical / Engineering professionals.

---

INPUTS:
- Product Documentation: See CLEANED TEXT below
- Job Role: Technical professionals including Backend Engineers, DevOps Engineers, QA Engineers, and Solutions Engineers
- Course End Goal: Learners will be proficient in configuring, executing, and managing data extraction processes using Oracle Fusion Business Intelligence Cloud Connector (BICC). They will be capable of integrating extracted data into external systems or data warehouses, ensuring data integrity, security, and compliance with organizational standards.

---

YOUR TASK:
Carefully read the product documentation provided. Then, using the Job Role and
Course End Goal as your north star, design a complete training course that is
laser-focused on getting the learner to that end goal. Every module and lesson
must connect back to achieving that goal.

The course depth and length should be flexible — let the complexity of the
documentation and the ambition of the end goal guide how many modules are needed.

---

COURSE OUTPUT FORMAT:

## 🏷️ Course Title
A clear, role-specific title (e.g., "Mastering [Product Name] for DevOps Engineers")

## 📋 Course Overview
- Target Audience: [Job Role]
- Estimated Duration: [Based on content complexity]
- Prerequisites: [Any assumed knowledge]

## 🎯 Course End Goal
State clearly and specifically what the learner will be able to DO, BUILD, or
ACHIEVE by the end of this course. Directly reflect the end goal provided,
expanded into a full, inspiring outcome statement.

## 🗺️ Learning Journey Summary
In 3-5 sentences, describe the journey the learner will go through — from where
they start to where they will be at the end.

---

## 📚 MODULE BREAKDOWN

For each module and its lessons, use the following structure:

---

### Module [Number]: [Module Name]
**How this module connects to the End Goal:** [One sentence explaining relevance]
**Duration:** [Estimated time]

#### Lesson [Number]: [Lesson Name]

**What We Teach in This Lesson:**
A clear paragraph explaining exactly what is covered in this lesson — the concepts,
features, or skills the learner will be walked through.

**Key Takeaways:**
- [Takeaway 1 — a specific thing the learner now knows or can do]
- [Takeaway 2]
- [Takeaway 3]

---

(Repeat the Lesson structure for all lessons within the module)
(Repeat the Module structure for all modules)

---

## 🎓 What You Can Do Now
Write a confident, motivating closing statement (3-5 sentences) telling the learner
exactly what they are now capable of doing with the product — directly tied to the
end goal they set out to achieve.

## ✅ End Goal Achievement Checklist
A 5-8 point checklist the learner can use to self-verify they have achieved
the course end goal. Each item should be a concrete, testable action.

- [ ] I can [action 1]
- [ ] I can [action 2]
- [ ] I can [action 3]
...

---

## 📝 Assessment Topics

Based on the course content above, identify the key topics that should be covered
in assessments for this course. Do NOT write the actual questions — only provide
the topic areas and a brief explanation of why each topic is important to test.

### Assessment Topic Areas

**Topic [Number]: [Topic Name]**
- Why test this: [1-2 sentences explaining what this topic validates]
- Suggested assessment type: [Multiple Choice / Scenario-Based / Practical Exercise / True or False / Fill in the Blank]
- Difficulty level: [Foundational / Intermediate / Advanced]

Include 5-10 assessment topic areas, ordered from foundational to advanced.

---

## 📣 Go-to-Market Section

Generate two pieces of ready-to-use marketing content to promote this course.

### 1. Social Media Post
Write a LinkedIn post (150-250 words) that:
- Opens with a bold hook
- Describes what the course covers and who it is for
- Highlights 2-3 specific skills or outcomes
- Ends with a clear call to action
- Includes 5-8 relevant hashtags
- Tone: Professional but energetic, like a knowledgeable practitioner

### 2. Newsletter Write-Up
Write a short newsletter feature (200-300 words) that includes:
- A compelling headline
- An opening paragraph addressing a real challenge the audience faces
- A "What You Will Learn" section with 4-6 bullet points written as benefits
- A closing paragraph with a call to action
- Tone: Authoritative, helpful, and customer-centric

---

IMPORTANT GUIDELINES:
- The Course End Goal is your compass — every lesson must serve it.
- Always tie content to what a Technical / Engineering professional needs to know and do.
- Use technical language appropriate for engineers.
- Keep takeaways concrete and action-oriented (use verbs like: configure, deploy, troubleshoot, integrate, optimize).
- Do not invent features or capabilities — only use what is in the documentation.
- For Go-to-Market: draw specific details from the course. Avoid vague phrases.

---

CLEANED TEXT FROM DOCUMENTATION:
{clean_text[:12000]}"""
                        }
                    ]
                )
                ai_response = chat_completion.choices[0].message.content

            except Exception as e:
                st.error(f"❌ AI error: {e}")
                st.stop()

        with st.spinner("Step 3/3 — Building your Word document..."):
            try:
                doc = Document()
                doc.add_heading('Course Design Document', 0)
                doc.add_heading('Source Content Summary', level=1)
                doc.add_paragraph(clean_text[:2000] + "..." if len(clean_text) > 2000 else clean_text)
                doc.add_heading('AI-Generated Course Design', level=1)
                doc.add_paragraph(ai_response)
                docx_buffer = io.BytesIO()
                doc.save(docx_buffer)
                docx_buffer.seek(0)

            except Exception as e:
                st.error(f"❌ Word doc error: {e}")
                st.stop()

        st.success("✅ Done! Your Course Design Document is ready.")
        st.markdown("### 📋 Generated Course Design")
        st.markdown(ai_response)

        original_name = os.path.splitext(uploaded_file.name)[0]
        st.download_button(
            label="📄 Download Word Document",
            data=docx_buffer,
            file_name=f"{original_name}_CourseDesign.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
else:
    st.info("📂 Please upload a PDF to proceed.")
